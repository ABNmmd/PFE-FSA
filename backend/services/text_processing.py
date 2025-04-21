import re
import spacy
import os
import io

# Optional imports with fallbacks
try:
    from langdetect import detect
except ImportError:
    print("Warning: langdetect not installed. Language detection disabled.")
    def detect(text):
        return 'fr'  # Default to French

# Document processing libraries - handle missing dependencies
pdf_support = True
docx_support = True
textract_support = True

try:
    import PyPDF2
except ImportError:
    print("Warning: PyPDF2 not installed. PDF processing will be limited.")
    pdf_support = False

try:
    import docx
except ImportError:
    print("Warning: python-docx not installed. DOCX processing will be limited.")
    docx_support = False

try:
    import textract
except ImportError:
    print("Warning: textract not installed. Advanced document processing will be limited.")
    textract_support = False

# Load spaCy model - use the French model since we're working with French documents
try:
    nlp = spacy.load('fr_core_news_sm')
    print("French language model loaded successfully.")
except:
    try:
        # Try to load a basic model if French is not available
        nlp = spacy.load('en_core_web_sm')
        print("Warning: French model not found, using English model instead.")
    except:
        print("Warning: No spaCy model found. Using basic text processing.")
        nlp = None

class TextProcessingService:
    """Service for processing text before plagiarism detection."""
    
    def __init__(self, debug=False):
        """Initialize the text processing service."""
        self.stopwords = set()
        self.debug = debug  # Debug flag to control print statements
        # Minimum meaningful chunk length (to avoid matching single words/punctuation)
        self.min_chunk_length = 25  # Characters
        # Minimum sentence length to consider (to avoid matching trivial sentences)
        self.min_sentence_length = 15  # Characters
    
    def preprocess_text(self, text, apply_lemmatization=False):
        """Clean and preprocess text, with optional lemmatization for TF-IDF."""
        if not text:
            return ""
            
        # Fix encoding issues first
        text = self._fix_encoding(text)
            
        # Convert to lowercase and remove extra whitespace
        text = text.lower()
        text = re.sub(r'\s+', ' ', text).strip()
        
        # Remove special characters, keeping only letters, numbers and basic punctuation
        text = re.sub(r'[^\w\s.,?!-]', '', text)
        
        # Apply lemmatization if requested (only used for TF-IDF)
        if apply_lemmatization and nlp:
            return self._lemmatize_text(text)
        
        return text
    
    def _fix_encoding(self, text):
        """Fix common encoding issues, especially with French text"""
        if not text:
            return ""
            
        # Handle surrogate characters that cause encoding errors
        try:
            # Replace or remove surrogate pairs and other problematic Unicode
            text = text.encode('utf-8', errors='surrogateescape').decode('utf-8', errors='replace')
        except:
            # If that fails, do a more aggressive replacement
            text = ''.join(c if ord(c) < 0x10000 else '\ufffd' for c in text)
        
        # Common encoding replacements for French
        replacements = {
            'Ã©': 'é', 'Ã¨': 'è', 'Ãª': 'ê', 'Ã«': 'ë',
            'Ã¢': 'â', 'Ã®': 'î', 'Ã´': 'ô', 'Ã»': 'û',
            'Ã¹': 'ù', 'Ã§': 'ç', 'Ã¯': 'ï', 'Ã¤': 'ä',
            'Ã¶': 'ö', 'Ã¼': 'ü', 'Ã': 'À', 'â\x80\x99': "'",
            'â\x80\x93': "-", 'â\x80\x94': "-", 'â\x80\x9c': '"', 'â\x80\x9d': '"'
        }
        
        for wrong, right in replacements.items():
            text = text.replace(wrong, right)
            
        return text
    
    def _lemmatize_text(self, text):
        """Lemmatize French text using spaCy"""
        if not text or not nlp:
            return text
            
        # Check if we have the French spaCy model
        is_french_model = False
        try:
            is_french_model = nlp.meta.get('lang', '') == 'fr'
            if self.debug and not is_french_model:
                print("[DEBUG] Warning: Using non-French spaCy model for lemmatization")
        except:
            pass
        
        try:
            # Process with spaCy
            doc = nlp(text)
            
            # Keep punctuation but lemmatize words
            lemmatized_tokens = []
            for token in doc:
                if token.is_punct or token.is_space:
                    lemmatized_tokens.append(token.text)
                else:
                    # Use lemma_ property which gives the base form
                    lemmatized_tokens.append(token.lemma_)
            
            result = ' '.join(lemmatized_tokens)
            
            if self.debug:
                print(f"[DEBUG] French lemmatization applied using spaCy")
                if len(text) > 200:
                    print(f"[DEBUG] Sample: '{text[:100]}...' → '{result[:100]}...'")
            
            return result
            
        except Exception as e:
            if self.debug:
                print(f"[DEBUG] Error during spaCy lemmatization: {str(e)}")
            return text  # Fall back to original text on error
    
    def split_into_sentences(self, text):
        """Split text into sentences using spaCy, filtering out trivial ones."""
        if not text:
            return []
            
        # Fix encoding issues first
        text = self._fix_encoding(text)
            
        # Try to detect language
        try:
            lang = detect(text[:100])  # Use first 100 chars for detection
        except:
            lang = 'fr'  # Default to French
            
        # Process with spaCy for better sentence segmentation
        if nlp:
            try:
                doc = nlp(text)
                # Filter out trivial sentences
                sentences = [sent.text.strip() for sent in doc.sents 
                            if sent.text.strip() and len(sent.text.strip()) >= self.min_sentence_length]
                
                if self.debug:
                    print(f"[DEBUG] Extracted {len(sentences)} meaningful sentences (min length: {self.min_sentence_length})")
                
                return sentences
            except Exception as e:
                if self.debug:
                    print(f"[DEBUG] Error splitting sentences with spaCy: {e}")
        
        # Fallback: regex-based sentence splitting
        raw_sentences = re.split(r'(?<=[.!?])\s+', text)
        # Filter out trivial sentences
        sentences = [s.strip() for s in raw_sentences if s.strip() and len(s.strip()) >= self.min_sentence_length]
        
        if self.debug:
            print(f"[DEBUG] Extracted {len(sentences)} meaningful sentences using fallback method")
            
        return sentences
    
    def chunk_document(self, text, chunk_size=200, overlap=50):
        """
        Chunk document into overlapping chunks.
        This is a simple chunking method that doesn't respect sentence boundaries.
        """
        if not text:
            return []
            
        # Fix encoding issues
        text = self._fix_encoding(text)
        
        # Clean text
        text = re.sub(r'\s+', ' ', text).strip()
        
        if self.debug:
            print(f"\n[DEBUG] CHUNKING DOCUMENT (size={chunk_size}, overlap={overlap})")
            print(f"[DEBUG] Original text length: {len(text)} characters")
        
        # Create chunks with overlap
        chunks = []
        for i in range(0, len(text), chunk_size - overlap):
            chunk = text[i:i + chunk_size]
            if len(chunk) >= chunk_size / 2:  # Only add if chunk is substantial
                chunks.append(chunk)
        
        # Filter out chunks that are too small
        chunks = [chunk for chunk in chunks if len(chunk) >= self.min_chunk_length]
        
        if self.debug:
            print(f"[DEBUG] Created {len(chunks)} chunks")
            for i, chunk in enumerate(chunks[:3]):  # Print first 3 chunks
                print(f"[DEBUG] Chunk {i}: {chunk[:100]}..." if len(chunk) > 100 else f"[DEBUG] Chunk {i}: {chunk}")
            if len(chunks) > 3:
                print(f"[DEBUG] ... and {len(chunks) - 3} more chunks")
                
        return chunks
    
    def chunk_document_semantic(self, text, min_chunk_size=80, max_chunk_size=500):
        """
        Create meaningful chunks from document text using spaCy.
        Respects sentence boundaries for better semantic comparison.
        Filters out trivial chunks.
        """
        if not text:
            return []
            
        # Clean text and fix encoding
        text = re.sub(r'\s+', ' ', text).strip()
        text = self._fix_encoding(text)
        
        if self.debug:
            print(f"\n[DEBUG] SEMANTIC CHUNKING (min={min_chunk_size}, max={max_chunk_size})")
            print(f"[DEBUG] Original text length: {len(text)} characters")
        
        # Process with spaCy for better sentence segmentation
        sentences = self.split_into_sentences(text)
        
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            # Skip very short sentences (likely not meaningful)
            if len(sentence.strip()) < self.min_sentence_length:
                if self.debug and sentence.strip():
                    print(f"[DEBUG] Skipping short sentence: {sentence}")
                continue
                
            # If adding this sentence would exceed max size, save current chunk and start new one
            if len(current_chunk) + len(sentence) <= max_chunk_size:
                current_chunk += " " + sentence if current_chunk else sentence
            else:
                if len(current_chunk) >= min_chunk_size:
                    chunks.append(current_chunk.strip())
                current_chunk = sentence
        
        # Add the last chunk if it's big enough
        if current_chunk and len(current_chunk) >= min_chunk_size:
            chunks.append(current_chunk.strip())
        
        # Final filter to ensure all chunks are meaningful
        chunks = [chunk for chunk in chunks if len(chunk) >= self.min_chunk_length]
        
        if self.debug:
            print(f"[DEBUG] Created {len(chunks)} semantic chunks")
            if chunks:
                for i, chunk in enumerate(chunks[:2]):
                    print(f"[DEBUG] Chunk {i}: {chunk[:100]}..." if len(chunk) > 100 else f"[DEBUG] Chunk {i}: {chunk}")
                if len(chunks) > 2:
                    print(f"[DEBUG] ... and {len(chunks) - 2} more chunks")
        
        return chunks
    
    def prepare_text(self, text, chunk_size=200, overlap=50):
        """
        Prepares text for TF-IDF similarity comparison by:
        1. Preprocessing (cleaning)
        2. Chunking
        3. Additional cleaning of chunks including lemmatization
        """
        # Fix encoding issues
        text = self._fix_encoding(text)
        
        # Preprocess the full text (no lemmatization yet for better chunking)
        cleaned_text = self.preprocess_text(text, apply_lemmatization=False)
        
        # Create chunks
        chunks = self.chunk_document(cleaned_text, chunk_size, overlap)
        
        # Additional cleaning of chunks WITH spaCy French lemmatization for TF-IDF
        cleaned_chunks = [self.preprocess_text(chunk, apply_lemmatization=True) for chunk in chunks]
        
        # Filter out empty chunks
        cleaned_chunks = [chunk for chunk in cleaned_chunks if chunk.strip()]
        
        if self.debug:
            print(f"[DEBUG] TF-IDF preparation with spaCy French lemmatization: {len(cleaned_chunks)} chunks")
        
        return cleaned_chunks
    
    def prepare_text_semantic(self, text):
        """
        Prepares text for semantic similarity comparison:
        1. Creates meaningful chunks respecting sentence boundaries
        2. Light preprocessing to maintain semantic meaning
        3. NO lemmatization as embeddings work better with original text
        """
        # Fix encoding issues
        text = self._fix_encoding(text)
        
        # Create semantic chunks
        chunks = self.chunk_document_semantic(text)
        
        # Light cleaning that maintains semantic meaning WITHOUT lemmatization
        cleaned_chunks = [chunk.strip() for chunk in chunks if chunk.strip()]
        
        if self.debug:
            print(f"[DEBUG] Embedding preparation without lemmatization: {len(cleaned_chunks)} chunks")
        
        return cleaned_chunks


def extract_text_content(file_obj, file_type, debug=False):
    """
    Extract text content from various file types with improved encoding handling.
    
    Args:
        file_obj: File-like object
        file_type: Type of file (pdf, txt, docx, etc.)
        debug: Whether to print debug information
    
    Returns:
        str: Extracted text content
    """
    if debug:
        print(f"\n[DEBUG] EXTRACTING TEXT FROM: {file_type}")
    
    try:
        if file_type == 'txt':
            # For text files, use utf-8 encoding with error handling
            file_obj.seek(0)
            return file_obj.read().decode('utf-8', errors='replace')
        
        elif file_type.lower() == 'pdf':
            # Fixed PDF extraction logic
            if pdf_support:
                try:
                    if debug:
                        print("[DEBUG] Using PyPDF2 for PDF extraction")
                    
                    # Always seek to beginning of file
                    file_obj.seek(0)
                    
                    reader = PyPDF2.PdfReader(file_obj)
                    if len(reader.pages) == 0:
                        if debug:
                            print("[DEBUG] PDF has no pages")
                        return "[Empty PDF document]"
                    
                    text = ""
                    for i, page in enumerate(reader.pages):
                        try:
                            page_text = page.extract_text() or ""  # Handle None
                            
                            # Clean problematic Unicode characters
                            page_text = _clean_unicode(page_text)
                            
                            text += page_text + "\n"
                            if debug and i < 1:
                                print(f"[DEBUG] Page {i+1} sample: {page_text[:100]}..." if page_text and len(page_text) > 100 else f"[DEBUG] Page {i+1} sample: {page_text or '(empty)'}")
                        except Exception as page_err:
                            if debug:
                                print(f"[DEBUG] Error extracting page {i+1}: {str(page_err)}")
                            text += f"[Error extracting page {i+1}]\n"
                    
                    # Final Unicode cleaning before returning
                    text = _clean_unicode(text)
                    
                    # Ensure we return a non-empty string
                    if not text.strip():
                        if debug:
                            print("[DEBUG] Extracted empty text from PDF")
                        return "[PDF contained no extractable text - may be scanned or image-based]"
                    
                    if debug:
                        print(f"[DEBUG] Successfully extracted {len(text)} characters from PDF")
                    return text
                    
                except Exception as pdf_err:
                    if debug:
                        print(f"[DEBUG] PyPDF2 extraction error: {str(pdf_err)}")
                    
                    # Return error message
                    return f"[PDF text extraction failed: {str(pdf_err)}]"
            else:
                if debug:
                    print("[DEBUG] PDF support not available - PyPDF2 not installed")
                return "[PDF TEXT EXTRACTION ERROR: PDF libraries not installed]"
        
        elif file_type == 'docx':
            # Use python-docx with better encoding handling
            import docx
            import io
            
            # Write the file_obj to a bytes IO object
            file_bytes = io.BytesIO()
            file_obj.seek(0)
            file_bytes.write(file_obj.read())
            file_bytes.seek(0)
            
            try:
                doc = docx.Document(file_bytes)
                # Extract text with proper encoding
                text = "\n".join([para.text for para in doc.paragraphs])
                
                # If text is empty, try extracting from runs to handle complex formatting
                if not text.strip():
                    text = "\n".join([
                        "\n".join([run.text for run in para.runs])
                        for para in doc.paragraphs
                    ])
                
                return text
            except Exception as e:
                if debug:
                    print(f"[DEBUG] Error extracting with python-docx: {str(e)}")
                
                # Fallback to using the docx2txt library
                try:
                    import docx2txt
                    file_obj.seek(0)
                    return docx2txt.process(file_obj)
                except Exception as e2:
                    if debug:
                        print(f"[DEBUG] Error with docx2txt fallback: {str(e2)}")
                    
                    # Ultimate fallback to extract what we can
                    return f"[Document text extraction partially failed: {str(e)}]"
        
        # Add other file type handlers as needed
        # ...existing code for other formats...
        
        else:
            if debug:
                print(f"[DEBUG] Unsupported file type: {file_type}")
            return f"[Unsupported file type: {file_type}]"
            
    except UnicodeDecodeError as ude:
        if debug:
            print(f"[DEBUG] ERROR extracting text: {str(ude)}")
            
        # Return a placeholder but don't completely fail
        # Try with latin-1 encoding as a fallback
        try:
            file_obj.seek(0)
            return file_obj.read().decode('latin-1', errors='replace')
        except:
            return f"[Document contains special characters that couldn't be decoded]"
            
    except Exception as e:
        if debug:
            print(f"[DEBUG] ERROR extracting text: {str(e)}")
        return f"[Document text extraction error: {type(e).__name__}]"

def _clean_unicode(text):
    """Clean problematic Unicode characters from text"""
    if not text:
        return ""
    
    try:
        # First attempt: Fix surrogate pairs
        cleaned = text.encode('utf-8', errors='surrogateescape').decode('utf-8', errors='replace')
        
        # Second cleaning: Remove any remaining problematic characters
        cleaned = ''.join(
            c for c in cleaned 
            if not (0xD800 <= ord(c) <= 0xDFFF)  # Remove surrogate code points
        )
        
        # Replace emoji and other special characters with placeholders
        emoji_pattern = re.compile(
            "["
            "\U0001F600-\U0001F64F"  # emoticons
            "\U0001F300-\U0001F5FF"  # symbols & pictographs
            "\U0001F680-\U0001F6FF"  # transport & map symbols
            "\U0001F700-\U0001F77F"  # alchemical symbols
            "\U0001F780-\U0001F7FF"  # Geometric Shapes
            "\U0001F800-\U0001F8FF"  # Supplemental Arrows-C
            "\U0001F900-\U0001F9FF"  # Supplemental Symbols and Pictographs
            "\U0001FA00-\U0001FA6F"  # Chess Symbols
            "\U0001FA70-\U0001FAFF"  # Symbols and Pictographs Extended-A
            "\U00002702-\U000027B0"  # Dingbats
            "\U000024C2-\U0001F251" 
            "]", flags=re.UNICODE
        )
        cleaned = emoji_pattern.sub(r'[emoji]', cleaned)
        
        return cleaned
    except Exception as e:
        # If all cleaning fails, do aggressive character filtering
        return ''.join(c for c in text if ord(c) < 0x10000)
